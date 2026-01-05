#!/usr/bin/env python3
"""
Convert notes files (.txt, .md, .docx) to standardized Markdown format

Processes notes from notes/raw/ directory:
- Reads .txt, .md, and .docx files
- Extracts metadata (timestamps, author if available)
- Converts to standardized Markdown format
- Saves to notes/ai/ directory
- Moves originals to notes/processed/

Supported formats:
- .txt - Plain text notes
- .md - Markdown notes
- .docx - Microsoft Word/OneNote exports
- .textbundle - Bear note bundles
- .html - Apple Notes HTML exports

Usage:
    python3 core/aiScripts/notesToMd/notes_to_md_converter.py
"""

import os
import re
import shutil
from pathlib import Path
from datetime import datetime
import sys
import zipfile
import json

# Import docx for OneNote support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Import html2text for Apple Notes HTML support
try:
    import html2text
    HTML2TEXT_AVAILABLE = True
except ImportError:
    HTML2TEXT_AVAILABLE = False

# Import logger
try:
    from ..logger import get_logger
    logger = get_logger('notes_converter')
except ImportError:
    # Fallback if running as standalone script
    sys.path.insert(0, str(Path(__file__).parent.parent))
    from logger import get_logger
    logger = get_logger('notes_converter')


def sanitize_filename(filename):
    """Sanitize filename to be safe for filesystem"""
    if not filename:
        return "unnamed_note"

    # Remove path separators and other dangerous characters
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)

    # Remove leading/trailing dots and spaces
    filename = filename.strip('. ')

    # Limit length to 200 characters
    if len(filename) > 200:
        name, ext = os.path.splitext(filename)
        filename = name[:200-len(ext)] + ext

    return filename if filename else "unnamed_note"


def extract_metadata(content, filename):
    """
    Extract metadata from note content

    Looks for common metadata patterns:
    - Date: YYYY-MM-DD or MM/DD/YYYY
    - Author: Author: Name or By: Name
    - Title: First line or filename

    Returns:
        dict: Metadata with keys: title, author, date, original_filename
    """
    metadata = {
        'title': None,
        'author': None,
        'date': None,
        'original_filename': filename
    }

    lines = content.split('\n')

    # Try to extract title from first non-empty line
    for line in lines:
        stripped = line.strip()
        if stripped and not stripped.startswith('#'):
            # Remove common markdown title markers
            metadata['title'] = stripped.lstrip('#').strip()
            break

    # If no title found, use filename without extension
    if not metadata['title']:
        metadata['title'] = Path(filename).stem

    # Look for author patterns
    author_pattern = r'(?:Author|By|From|Written by):\s*([^\n]+)'
    author_match = re.search(author_pattern, content, re.IGNORECASE)
    if author_match:
        metadata['author'] = author_match.group(1).strip()

    # Look for date patterns
    # Pattern 1: YYYY-MM-DD
    date_pattern1 = r'(\d{4}-\d{2}-\d{2})'
    # Pattern 2: MM/DD/YYYY
    date_pattern2 = r'(\d{1,2}/\d{1,2}/\d{4})'
    # Pattern 3: Month DD, YYYY
    date_pattern3 = r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4})'

    for pattern in [date_pattern1, date_pattern2, date_pattern3]:
        date_match = re.search(pattern, content, re.IGNORECASE)
        if date_match:
            metadata['date'] = date_match.group(1)
            break

    # If no date found, use file modification time
    if not metadata['date']:
        try:
            file_path = Path('notes/raw') / filename
            if file_path.exists():
                mtime = file_path.stat().st_mtime
                metadata['date'] = datetime.fromtimestamp(mtime).strftime('%Y-%m-%d')
        except Exception as e:
            logger.warning(f"Could not extract file modification time: {e}")
            metadata['date'] = datetime.now().strftime('%Y-%m-%d')

    return metadata


def convert_note_to_markdown(content, metadata):
    """
    Convert note content to standardized Markdown format

    Args:
        content: Original note content
        metadata: Extracted metadata dict

    Returns:
        str: Formatted Markdown content
    """
    markdown = []

    # Add metadata header
    markdown.append("---")
    markdown.append(f"# Note: {metadata['title']}")
    markdown.append("")

    if metadata.get('author'):
        markdown.append(f"**Author**: {metadata['author']}")

    if metadata.get('date'):
        markdown.append(f"**Date**: {metadata['date']}")

    markdown.append(f"**Source**: {metadata['original_filename']}")
    markdown.append("---")
    markdown.append("")

    # Add content
    # Clean up excessive whitespace
    cleaned_content = re.sub(r'\n{3,}', '\n\n', content.strip())
    markdown.append(cleaned_content)

    return '\n'.join(markdown)


def parse_docx(source_path):
    """
    Parse Microsoft Word/OneNote .docx file

    Args:
        source_path: Path to .docx file

    Returns:
        str: Extracted text content with preserved structure
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for .docx support. Install with: pip install python-docx")

    doc = Document(str(source_path))
    content_lines = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            # Preserve paragraph breaks
            if content_lines and content_lines[-1]:
                content_lines.append('')
            continue

        # Detect heading style
        style_name = paragraph.style.name.lower()
        
        if 'heading 1' in style_name or 'title' in style_name:
            content_lines.append(f"# {text}")
        elif 'heading 2' in style_name:
            content_lines.append(f"## {text}")
        elif 'heading 3' in style_name:
            content_lines.append(f"### {text}")
        elif 'heading 4' in style_name:
            content_lines.append(f"#### {text}")
        else:
            content_lines.append(text)

    # Extract text from tables
    if doc.tables:
        content_lines.append('')
        content_lines.append('---')
        content_lines.append('**Tables:**')
        content_lines.append('')
        
        for table_idx, table in enumerate(doc.tables, 1):
            content_lines.append(f'Table {table_idx}:')
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    content_lines.append(row_text)
            content_lines.append('')

    return '\n'.join(content_lines)


def parse_textbundle(source_path):
    """
    Parse Bear .textbundle format

    .textbundle can be either a directory or ZIP containing:
    - text.md (or text.txt) - main content
    - info.json - metadata (creation date, tags, etc.)
    - assets/ - embedded images (optional)

    Args:
        source_path: Path to .textbundle file or directory

    Returns:
        str: Extracted markdown content
    """
    content_lines = []
    metadata = {}

    # Check if it's a ZIP file or directory
    if source_path.is_file():
        # It's a ZIP file
        try:
            with zipfile.ZipFile(source_path, 'r') as zip_ref:
                # Look for info.json
                if 'info.json' in zip_ref.namelist():
                    try:
                        info_data = json.loads(zip_ref.read('info.json'))
                        metadata = info_data
                    except json.JSONDecodeError:
                        logger.warning(f"Could not parse info.json in {source_path.name}")

                # Look for text content (text.md or text.txt)
                text_file = None
                if 'text.md' in zip_ref.namelist():
                    text_file = 'text.md'
                elif 'text.txt' in zip_ref.namelist():
                    text_file = 'text.txt'

                if text_file:
                    content = zip_ref.read(text_file).decode('utf-8')
                    content_lines.append(content)
                else:
                    logger.warning(f"No text.md or text.txt found in {source_path.name}")
        except zipfile.BadZipFile:
            logger.error(f"Invalid ZIP file: {source_path.name}")
            return ""

    elif source_path.is_dir():
        # It's a directory
        info_path = source_path / 'info.json'
        if info_path.exists():
            try:
                with open(info_path, 'r', encoding='utf-8') as f:
                    metadata = json.load(f)
            except json.JSONDecodeError:
                logger.warning(f"Could not parse info.json in {source_path.name}")

        # Look for text content
        text_md = source_path / 'text.md'
        text_txt = source_path / 'text.txt'

        if text_md.exists():
            with open(text_md, 'r', encoding='utf-8') as f:
                content_lines.append(f.read())
        elif text_txt.exists():
            with open(text_txt, 'r', encoding='utf-8') as f:
                content_lines.append(f.read())
        else:
            logger.warning(f"No text.md or text.txt found in {source_path.name}")

    # Add metadata as comments if available
    if metadata:
        meta_header = []
        if 'creatorIdentifier' in metadata:
            meta_header.append(f"<!-- Creator: {metadata['creatorIdentifier']} -->")
        if 'version' in metadata:
            meta_header.append(f"<!-- Format Version: {metadata['version']} -->")
        if 'type' in metadata:
            meta_header.append(f"<!-- Type: {metadata['type']} -->")

        if meta_header:
            return '\n'.join(meta_header) + '\n\n' + '\n'.join(content_lines)

    return '\n'.join(content_lines)


def parse_html(source_path):
    """
    Parse Apple Notes HTML export

    Args:
        source_path: Path to .html file

    Returns:
        str: Markdown content converted from HTML
    """
    if not HTML2TEXT_AVAILABLE:
        raise ImportError("html2text is required for HTML support. Install with: pip install html2text")

    # Read HTML content
    with open(source_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    # Configure html2text for clean markdown conversion
    h = html2text.HTML2Text()
    h.ignore_links = False
    h.ignore_images = False
    h.ignore_emphasis = False
    h.body_width = 0  # Don't wrap lines
    h.unicode_snob = True  # Use unicode characters
    h.skip_internal_links = True
    h.ignore_mailto_links = False
    h.protect_links = True
    h.mark_code = True

    # Convert HTML to markdown
    markdown_content = h.handle(html_content)

    # Clean up excessive newlines (more than 2 consecutive)
    markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content)

    return markdown_content.strip()


def detect_format(file_path):
    """
    Detect file format from extension or directory name

    Args:
        file_path: Path to file or directory

    Returns:
        str: File format ('txt', 'md', 'docx', 'textbundle', 'html', 'unknown')
    """
    # Check if it's a .textbundle (can be file or directory)
    if file_path.suffix.lower() == '.textbundle':
        return 'textbundle'

    # Check if it's a directory with .textbundle extension
    if file_path.is_dir() and file_path.name.endswith('.textbundle'):
        return 'textbundle'

    suffix = file_path.suffix.lower()

    if suffix == '.txt':
        return 'txt'
    elif suffix == '.md':
        return 'md'
    elif suffix == '.docx':
        return 'docx'
    elif suffix == '.html':
        return 'html'
    else:
        return 'unknown'


def process_notes_file(source_path, raw_dir, ai_dir, processed_dir):
    """
    Process a single notes file

    Args:
        source_path: Path to source file
        raw_dir: Path to raw directory
        ai_dir: Path to AI directory
        processed_dir: Path to processed directory

    Returns:
        bool: True if successful, False otherwise
    """
    try:
        filename = source_path.name
        logger.info(f"Processing: {filename}")

        # Detect format and read content
        file_format = detect_format(source_path)

        if file_format == 'html':
            if not HTML2TEXT_AVAILABLE:
                logger.error(f"Skipping {filename}: html2text not available. Install with: pip install html2text")
                return False
            content = parse_html(source_path)
        elif file_format == 'textbundle':
            content = parse_textbundle(source_path)
        elif file_format == 'docx':
            if not DOCX_AVAILABLE:
                logger.error(f"Skipping {filename}: python-docx not available. Install with: pip install python-docx")
                return False
            content = parse_docx(source_path)
        elif file_format in ('txt', 'md'):
            # Read text/markdown files directly
            try:
                with open(source_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except UnicodeDecodeError:
                # Try with latin-1 encoding as fallback
                logger.warning(f"UTF-8 decode failed for {filename}, trying latin-1")
                with open(source_path, 'r', encoding='latin-1') as f:
                    content = f.read()
        else:
            logger.warning(f"Skipping {filename}: Unsupported format {source_path.suffix}")
            return False

        if not content.strip():
            logger.warning(f"Empty file: {filename}")
            return False

        # Extract metadata
        metadata = extract_metadata(content, filename)
        logger.debug(f"Extracted metadata: {metadata}")

        # Convert to markdown
        markdown_content = convert_note_to_markdown(content, metadata)

        # Generate output filename
        base_name = Path(filename).stem
        safe_name = sanitize_filename(base_name)
        output_filename = f"{safe_name}.md"

        # Save to AI directory
        output_path = ai_dir / output_filename
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)

        logger.info(f"Saved to: {output_path}")

        # Move original to processed directory
        processed_path = processed_dir / filename
        shutil.move(str(source_path), str(processed_path))
        logger.info(f"Moved original to: {processed_path}")

        return True

    except Exception as e:
        logger.error(f"Error processing {source_path.name}: {e}", exc_info=True)
        return False


def main():
    """Main conversion workflow"""
    logger.info("Starting notes to Markdown conversion")

    # Define directory paths
    raw_dir = Path('notes/raw')
    ai_dir = Path('notes/ai')
    processed_dir = Path('notes/processed')

    # Create directories if they don't exist
    for directory in [raw_dir, ai_dir, processed_dir]:
        directory.mkdir(parents=True, exist_ok=True)
        logger.debug(f"Ensured directory exists: {directory}")

# Find all notes files in raw directory (.txt, .md, .docx, .textbundle, .html)
    notes_files = (
        list(raw_dir.glob('*.txt')) +
        list(raw_dir.glob('*.md')) +
        list(raw_dir.glob('*.docx')) +
        list(raw_dir.glob('*.textbundle')) +
        list(raw_dir.glob('*.html'))
    )

    if not notes_files:
        logger.info("No notes files found in notes/raw/")
        logger.info("Place .txt, .md, .docx, .textbundle, or .html files in notes/raw/ to process them")
        return 0

    logger.info(f"Found {len(notes_files)} notes file(s) to process")

    # Process each file
    success_count = 0
    fail_count = 0

    for notes_file in notes_files:
        if process_notes_file(notes_file, raw_dir, ai_dir, processed_dir):
            success_count += 1
        else:
            fail_count += 1

    # Summary
    logger.info("=" * 60)
    logger.info("Conversion Summary:")
    logger.info(f"  Successfully converted: {success_count}")
    logger.info(f"  Failed: {fail_count}")
    logger.info(f"  Output directory: {ai_dir.absolute()}")
    logger.info(f"  Processed files moved to: {processed_dir.absolute()}")
    logger.info("=" * 60)

    return 0 if fail_count == 0 else 1


if __name__ == '__main__':
    sys.exit(main())
